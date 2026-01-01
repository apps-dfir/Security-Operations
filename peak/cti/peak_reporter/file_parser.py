from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional
import hashlib
import os
import re

# NOTE:
# - SourceContent is a structured container.
# - load_source_as_text() MUST return a string (not SourceContent),
#   because callers pass it into MITRE extraction and markdown generation.


# ---------------------------------------------------------------------------
# Types
# ---------------------------------------------------------------------------

@dataclass
class SourceContent:
    source: str
    text: str
    title: str = ""
    ext: str = ""
    kind: str = ""  # "url" | "file"
    meta: dict = field(default_factory=dict)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_URL_RE = re.compile(r"^https?://", flags=re.I)


def is_url(s: str) -> bool:
    return bool(_URL_RE.match((s or "").strip()))


def guess_ext_from_url(url: str) -> str:
    u = (url or "").strip()
    u = u.split("#", 1)[0].split("?", 1)[0]
    return Path(u).suffix.lower()


def _sha256_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()


def _safe_mkdir(p: Path) -> None:
    p.mkdir(parents=True, exist_ok=True)


def _read_text_file(p: Path) -> str:
    # tolerant read
    try:
        return p.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return p.read_text(encoding="utf-8", errors="replace")


def _normalize_repo_relative_path(raw: str) -> str:
    """
    Accept repo-relative paths like:
      - peak/cti/inputs/foo.pdf
    Normalize to:
      - inputs/foo.pdf
    """
    s = (raw or "").strip().replace("\\", "/")
    if s.startswith("peak/cti/"):
        s = s[len("peak/cti/") :]
    return s


def _resolve_local_input(input_value: str, inputs_dir: Path) -> Path:
    """
    Resolve:
      - inputs/foo.pdf
      - peak/cti/inputs/foo.pdf
      - foo.pdf (assumed to be under inputs_dir)
      - absolute paths
      - other relative paths (as-is)
    """
    raw = _normalize_repo_relative_path(input_value)
    p = Path(raw)

    if p.is_absolute():
        return p

    # If it exists relative to CWD, use it
    if p.exists():
        return p

    # If user passed inputs/<file>, try under inputs_dir as well (covers odd CWDs)
    if raw.startswith("inputs/"):
        candidate = inputs_dir / Path(raw).name
        if candidate.exists():
            return candidate

    # If they gave just a filename, assume inputs_dir/<name>
    candidate = inputs_dir / p.name
    return candidate


# ---------------------------------------------------------------------------
# Extraction implementations
# ---------------------------------------------------------------------------

def _extract_pdf_text_docling(pdf_path: Path, cache_dir: Path) -> Optional[str]:
    """
    Docling-first extraction.
    Returns text on success, None if docling isn't available or fails.
    """
    try:
        # Docling APIs have moved a bit across versions.
        # We try a couple of known import paths.
        try:
            from docling.document_converter import DocumentConverter  # type: ignore
        except Exception:
            from docling import DocumentConverter  # type: ignore

        _safe_mkdir(cache_dir)

        # Convert
        converter = DocumentConverter()
        result = converter.convert(str(pdf_path))

        # Try common result/document access patterns
        doc = None
        for attr in ("document", "doc", "output", "result"):
            if hasattr(result, attr):
                doc = getattr(result, attr)
                break
        if doc is None:
            doc = result  # sometimes result itself is the document-like object

        # Try common export methods
        for fn in ("export_to_text", "to_text", "get_text", "text"):
            if hasattr(doc, fn):
                v = getattr(doc, fn)
                txt = v() if callable(v) else v
                if isinstance(txt, str) and txt.strip():
                    return txt

        for fn in ("export_to_markdown", "to_markdown", "markdown"):
            if hasattr(doc, fn):
                v = getattr(doc, fn)
                md = v() if callable(v) else v
                if isinstance(md, str) and md.strip():
                    # Use markdown as extracted text if that’s what we can get
                    return md

        # If nothing matched, return a string representation as last resort
        s = str(doc)
        return s if s.strip() else None

    except Exception:
        return None


def _extract_pdf_text_pdfplumber(pdf_path: Path) -> str:
    import pdfplumber  # type: ignore

    chunks: list[str] = []
    with pdfplumber.open(str(pdf_path)) as pdf:
        for page in pdf.pages:
            try:
                t = page.extract_text() or ""
            except Exception:
                t = ""
            if t.strip():
                chunks.append(t)
    return "\n\n".join(chunks).strip()


def _extract_docx_text(docx_path: Path) -> str:
    from docx import Document  # type: ignore

    doc = Document(str(docx_path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    return "\n".join(parts).strip()


def _extract_html_text_from_url(url: str, timeout_s: int = 60) -> SourceContent:
    import requests  # type: ignore
    from bs4 import BeautifulSoup  # type: ignore
    from readability import Document  # type: ignore

    r = requests.get(url, timeout=timeout_s, headers={"User-Agent": "PEAK-CTI/1.0"})
    r.raise_for_status()

    # Best-effort title/text extraction
    title = ""
    text = ""

    try:
        doc = Document(r.text)
        title = (doc.short_title() or "").strip()
        html = doc.summary(html_partial=True)
        soup = BeautifulSoup(html, "html.parser")
        text = soup.get_text("\n", strip=True)
    except Exception:
        soup = BeautifulSoup(r.text, "html.parser")
        title = (soup.title.get_text(strip=True) if soup.title else "") or ""
        text = soup.get_text("\n", strip=True)

    return SourceContent(
        source=url,
        title=title,
        text=text,
        ext=guess_ext_from_url(url),
        kind="url",
        meta={"content_type": r.headers.get("content-type", "")},
    )


def _download_url_to_cache(url: str, cache_dir: Path, timeout_s: int = 60) -> Path:
    import requests  # type: ignore

    _safe_mkdir(cache_dir)

    r = requests.get(url, timeout=timeout_s, headers={"User-Agent": "PEAK-CTI/1.0"}, stream=True)
    r.raise_for_status()

    data = r.content
    ext = guess_ext_from_url(url) or ""
    digest = _sha256_bytes(data)[:16]
    name = f"dl_{digest}{ext or ''}"
    out = cache_dir / name
    out.write_bytes(data)
    return out


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def load_source_as_text(
    source: str,
    inputs_dir: Path,
    cache_dir: Path,
    timeout_s: int = 60,
) -> str:
    """
    IMPORTANT: Returns a STRING (not SourceContent).

    - If source is a URL:
        - HTML gets extracted via readability/bs4
        - PDFs get downloaded and parsed docling-first
    - If source is a file path:
        - PDFs: docling-first, then pdfplumber
        - DOCX: python-docx
        - TXT/MD: plain text read
    """
    sc = load_source(
        source=source,
        inputs_dir=inputs_dir,
        cache_dir=cache_dir,
        timeout_s=timeout_s,
    )
    return sc.text or ""


def load_source(
    source: str,
    inputs_dir: Path,
    cache_dir: Path,
    timeout_s: int = 60,
) -> SourceContent:
    """
    Structured loader that returns SourceContent.
    Use load_source_as_text() when you only need the text.
    """
    s = (source or "").strip()

    if is_url(s):
        ext = guess_ext_from_url(s)

        # If it's a PDF (or looks like one), download and parse as PDF
        if ext == ".pdf" or re.search(r"\.pdf($|\?|#)", s, flags=re.I):
            pdf_path = _download_url_to_cache(s, cache_dir=cache_dir, timeout_s=timeout_s)
            txt = _extract_pdf_text_docling(pdf_path, cache_dir=cache_dir)
            if txt is None:
                txt = _extract_pdf_text_pdfplumber(pdf_path)
            return SourceContent(
                source=s,
                title=Path(s.split("?", 1)[0].split("#", 1)[0]).name,
                text=txt or "",
                ext=".pdf",
                kind="url",
                meta={"downloaded_to": str(pdf_path)},
            )

        # Default: treat as HTML/article
        return _extract_html_text_from_url(s, timeout_s=timeout_s)

    # Local file input
    p = _resolve_local_input(s, inputs_dir=inputs_dir)
    if not p.exists():
        raise FileNotFoundError(f"Input file not found: {s} (resolved to {p})")

    # Validate file size
    file_size = p.stat().st_size
    if file_size == 0:
        raise ValueError(f"File is empty (0 bytes): {p}")
    if file_size < 100:
        # Check if it's a Git LFS pointer file
        try:
            content = p.read_text(encoding='utf-8', errors='ignore')[:200]
            if content.startswith('version https://git-lfs.github.com'):
                raise ValueError(
                    f"File appears to be a Git LFS pointer (not the actual file). "
                    f"Run 'git lfs pull' to download the actual content: {p}"
                )
        except Exception:
            pass
        raise ValueError(f"File is suspiciously small ({file_size} bytes): {p}")

    ext = p.suffix.lower()
    title = p.name
    text = ""

    if ext == ".pdf":
        txt = _extract_pdf_text_docling(p, cache_dir=cache_dir)
        if txt is None:
            txt = _extract_pdf_text_pdfplumber(p)
        text = txt or ""

    elif ext == ".docx":
        text = _extract_docx_text(p)

    elif ext in (".txt", ".md"):
        text = _read_text_file(p)

    else:
        # best effort for unknown text-ish files
        text = _read_text_file(p)

    return SourceContent(
        source=str(p),
        title=title,
        text=text,
        ext=ext,
        kind="file",
        meta={"resolved_path": str(p)},
    )
